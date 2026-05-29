#!/usr/bin/env python3
"""Generate bilingual thesis documentation Word file: 解释项目.docx"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "docs" / "解释项目.docx"


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h1(doc, en, zh):
    doc.add_heading(f"{en} / {zh}", level=1)


def add_h2(doc, en, zh):
    doc.add_heading(f"{en} / {zh}", level=2)


def add_h3(doc, en, zh):
    doc.add_heading(f"{en} / {zh}", level=3)


def add_para(doc, en, zh=""):
    if zh:
        doc.add_paragraph(f"EN: {en}")
        p = doc.add_paragraph(f"ZH: {zh}")
    else:
        p = doc.add_paragraph(en)
    for run in p.runs:
        run.font.size = Pt(11)


def add_bullet(doc, en, zh):
    doc.add_paragraph(f"• EN: {en}", style="List Bullet")
    doc.add_paragraph(f"  ZH: {zh}", style="List Bullet")


def add_code(doc, code):
    p = doc.add_paragraph(code)
    p.style = "No Spacing"
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    add_title(doc, "解释项目 / Project Explanation Document")
    add_para(
        doc,
        "Full-Stack Study System — Design and Development",
        "全栈学习系统 — 设计与开发",
    )
    add_para(doc, "Author 作者: Chen Junshuo 陈骏烁", "")
    add_para(
        doc,
        "Thesis prototype & technical development guide (bilingual).",
        "论文原型与技术开发说明（中英对照）。",
    )
    doc.add_page_break()

    # --- 1 ---
    add_h1(doc, "1. Project Overview", "1. 项目概述")
    add_para(
        doc,
        "This project is a mini full-stack web application similar to a university study/course registration system.",
        "本项目是一个类似高校教务选课系统的迷你全栈 Web 应用。",
    )
    add_para(
        doc,
        "Students can register, log in, browse courses, enroll, drop courses, and view grades. Teachers can create, edit, delete courses, manage enrolled students, assign grades, and remove students from a course.",
        "学生可以注册、登录、浏览课程、选课、退课、查看成绩。教师可以创建、编辑、删除课程，管理选课学生，录入成绩，并将学生从课程中移除。",
    )
    add_bullet(
        doc,
        "Frontend: HTML5, CSS3, JavaScript",
        "前端：HTML5、CSS3、JavaScript",
    )
    add_bullet(doc, "Backend: Python 3 + Flask 3", "后端：Python 3 + Flask 3")
    add_bullet(doc, "Database: SQLite (file study_system.db)", "数据库：SQLite 文件 study_system.db")
    add_bullet(
        doc,
        "Repository: https://github.com/qwerscd/full-stack-study-system",
        "代码仓库：https://github.com/qwerscd/full-stack-study-system",
    )

    add_h2(doc, "1.1 Why no models.py?", "1.1 为什么没有 models.py？")
    add_para(
        doc,
        "Many Flask tutorials use SQLAlchemy with a models.py file. This thesis project intentionally uses database.py with raw SQL for simplicity and to demonstrate relational database design directly.",
        "许多 Flask 教程使用 SQLAlchemy 和 models.py。本论文项目有意使用 database.py 直接写 SQL，结构更简单，也便于在论文中展示关系型数据库设计。",
    )
    add_para(
        doc,
        "database.py = connection + table creation + seed data (equivalent to models + migrations in larger projects).",
        "database.py = 数据库连接 + 建表 + 演示数据（在大项目中相当于 models + 迁移）。",
    )

    # --- 2 Development process ---
    add_h1(doc, "2. Development Process (Thesis Timeline)", "2. 开发过程（论文时间线）")
    steps = [
        ("Month 1: Requirements", "第1月：需求分析", "Define users, roles, enroll rules, max capacity, duplicate prevention.", "确定用户、角色、选课规则、人数上限、防重复选课。"),
        ("Month 2: Design", "第2月：系统设计", "Design ER diagram, tables (users, courses, enrollments), page flow.", "设计 ER 图、三张表、页面流程。"),
        ("Month 3: Theory writing", "第3月：理论部分", "Write chapters on full-stack, DB, auth, UX.", "撰写全栈、数据库、认证、UX 理论章节。"),
        ("Month 4–5: Implementation", "第4–5月：编码实现", "Build Flask routes, templates, CSS, JS, SQLite.", "实现 Flask 路由、模板、样式、脚本、SQLite。"),
        ("Month 6: Testing", "第6月：测试", "Test login, enroll, full course, drop, teacher CRUD, permissions.", "测试登录、选课、满员、退课、教师增删改、权限。"),
        ("Month 7: Thesis completion", "第7月：论文定稿", "Document results, screenshots, GitHub upload.", "整理结果、截图、上传 GitHub。"),
    ]
    for en_t, zh_t, en_d, zh_d in steps:
        add_h3(doc, en_t, zh_t)
        add_para(doc, en_d, zh_d)

    # --- 3 Architecture ---
    add_h1(doc, "3. System Architecture", "3. 系统架构")
    add_code(
        doc,
        """Browser (HTML/CSS/JS)
    ↓ HTTP GET/POST
Flask app.py (routes, session, flash)
    ↓ SQL
database.py → SQLite (study_system.db)""",
    )
    add_para(
        doc,
        "This is a classic three-tier architecture: presentation (templates), application (Flask), data (SQLite).",
        "这是经典三层架构：表示层（模板）、应用层（Flask）、数据层（SQLite）。",
    )

    add_h2(doc, "3.1 Request lifecycle", "3.1 一次请求的生命周期")
    add_bullet(doc, "User clicks a link or submits a form.", "用户点击链接或提交表单。")
    add_bullet(doc, "Browser sends HTTP request to Flask (e.g. POST /student/enroll/1).", "浏览器向 Flask 发送 HTTP 请求。")
    add_bullet(doc, "@login_required checks session and role.", "装饰器 login_required 检查会话与角色。")
    add_bullet(doc, "Route function runs SQL via get_db() or db_session().", "路由函数通过 get_db 或 db_session 执行 SQL。")
    add_bullet(doc, "render_template() returns HTML, or redirect() sends user elsewhere.", "render_template 返回 HTML，或 redirect 跳转。")
    add_bullet(doc, "teardown closes database connection.", "请求结束时 teardown 关闭数据库连接。")

    # --- 4 File structure ---
    add_h1(doc, "4. Project File Structure", "4. 项目文件结构")
    add_code(
        doc,
        """full-stack-study-system/
├── app.py                 # Flask routes (main logic)
├── database.py            # SQLite schema & seed (not models.py)
├── requirements.txt       # Dependencies
├── static/css/style.css   # UI styles
├── static/js/app.js       # Confirm & flash UX
└── templates/             # Jinja2 HTML pages
    ├── base.html
    ├── login.html, register.html
    ├── student/ (courses, my_courses, grades)
    └── teacher/ (courses, course_form, students)""",
    )

    # --- 5 Database ---
    add_h1(doc, "5. Database Design & SQL Syntax", "5. 数据库设计与 SQL 语法")
    add_h2(doc, "5.1 Table: users", "5.1 表：users")
    add_para(
        doc,
        "Stores login accounts. role is CHECK-constrained to 'student' or 'teacher'. password_hash never stores plain text.",
        "存储登录账号。role 用 CHECK 限制为 student 或 teacher。password_hash 不存明文密码。",
    )
    add_code(
        doc,
        """CREATE TABLE users (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  username TEXT NOT NULL UNIQUE,
  email TEXT NOT NULL UNIQUE,
  password_hash TEXT NOT NULL,
  role TEXT NOT NULL CHECK (role IN ('student', 'teacher')),
  created_at TEXT DEFAULT (datetime('now'))
);""",
    )

    add_h2(doc, "5.2 Table: courses", "5.2 表：courses")
    add_para(
        doc,
        "Each course belongs to one teacher (teacher_id FK). max_capacity enforces enrollment limit in app logic.",
        "每门课属于一名教师（外键 teacher_id）。max_capacity 在应用层与选课逻辑配合实现人数上限。",
    )

    add_h2(doc, "5.3 Table: enrollments", "5.3 表：enrollments")
    add_para(
        doc,
        "Links student to course. UNIQUE(student_id, course_id) prevents duplicate enrollment at database level.",
        "连接学生与课程。UNIQUE(student_id, course_id) 在数据库层防止重复选课。",
    )
    add_para(
        doc,
        "ON DELETE CASCADE: deleting a user or course automatically removes related enrollments.",
        "ON DELETE CASCADE：删除用户或课程时，相关选课记录自动删除。",
    )

    add_h2(doc, "5.4 Python database syntax used", "5.4 使用的 Python 数据库语法")
    add_bullet(doc, "sqlite3.connect(path) — open file database", "sqlite3.connect — 打开文件数据库")
    add_bullet(doc, "conn.row_factory = sqlite3.Row — access columns by name", "row_factory — 按列名访问结果")
    add_bullet(doc, "conn.execute(sql, (param,)) — parameterized query (prevents SQL injection)", "参数化查询 ? — 防止 SQL 注入")
    add_bullet(doc, "@contextmanager + with db_session() — commit/rollback/close", "上下文管理器 — 自动提交/回滚/关闭")
    add_bullet(doc, "PRAGMA foreign_keys = ON — enable foreign keys in SQLite", "启用 SQLite 外键约束")

    doc.add_page_break()

    # --- 6 Flask syntax ---
    add_h1(doc, "6. Flask Syntax & Core Concepts", "6. Flask 语法与核心概念")
    concepts = [
        ("Flask(__name__)", "创建应用对象", "Creates the application object.", "创建应用对象。"),
        ("@app.route('/path', methods=['GET','POST'])", "路由装饰器", "Maps URL to Python function.", "将 URL 映射到 Python 函数。"),
        ("render_template('x.html', var=value)", "渲染模板", "Renders Jinja2 HTML with variables.", "用变量渲染 Jinja2 模板。"),
        ("request.form.get('field')", "读取表单", "Reads POST form field from browser.", "读取浏览器提交的表单字段。"),
        ("session['user_id']", "会话", "Stores logged-in user in signed cookie.", "在签名 Cookie 中保存登录状态。"),
        ("flash('message', 'category')", "一次性提示", "Shows message once on next page.", "在下页显示一次后消失的消息。"),
        ("redirect(url_for('login'))", "重定向", "Sends HTTP redirect to another route.", "HTTP 重定向到另一路由。"),
        ("g.current_user", "请求全局对象", "Stores per-request data (current user).", "存放单次请求的数据（当前用户）。"),
        ("@app.teardown_appcontext", "清理", "Runs after request (close DB).", "请求结束后执行（关闭数据库）。"),
    ]
    for syn, zh_s, en_d, zh_d in concepts:
        add_h3(doc, syn, zh_s)
        add_para(doc, en_d, zh_d)

    # --- 7 Auth ---
    add_h1(doc, "7. Authentication & Authorization", "7. 认证与权限控制")
    add_h2(doc, "7.1 Registration", "7.1 注册")
    add_para(
        doc,
        "generate_password_hash(password) from Werkzeug hashes the password before INSERT. Never store plain passwords.",
        "注册时用 Werkzeug 的 generate_password_hash 哈希后再 INSERT，绝不存明文。",
    )
    add_h2(doc, "7.2 Login", "7.2 登录")
    add_para(
        doc,
        "check_password_hash(stored_hash, typed_password) compares securely. On success: session['user_id'], session['role'].",
        "check_password_hash 安全比对密码。成功后将 user_id、role 写入 session。",
    )
    add_h2(doc, "7.3 login_required decorator", "7.3 login_required 装饰器")
    add_para(
        doc,
        "A decorator wraps route functions. If not logged in → redirect login. If role='student' but user is teacher → flash error and redirect.",
        "装饰器包裹路由：未登录→跳转登录；角色不符→提示无权限并跳转。",
    )
    add_code(
        doc,
        """@login_required(role="student")
def student_courses():
    ...  # only students can enter""",
    )
    add_h2(doc, "7.4 Ownership checks (teachers)", "7.4 所有权校验（教师）")
    add_para(
        doc,
        "Every teacher action includes AND teacher_id = g.current_user['id'] in SQL so teachers cannot edit/delete others' courses.",
        "教师操作 SQL 均带 teacher_id = 当前用户 id，无法修改他人课程。",
    )

    # --- 8 Business rules ---
    add_h1(doc, "8. Business Rules & Implementation", "8. 业务规则与实现")
    rules = [
        ("Max enrollment", "选课人数上限",
         "COUNT enrollments for course_id; if count >= max_capacity, flash error and return.",
         "统计该课选课数；若 >= max_capacity，提示满员并拒绝。"),
        ("No duplicate enrollment", "禁止重复选课",
         "SELECT existing row; UNIQUE constraint in DB as backup.",
         "先查询是否已选；数据库 UNIQUE 作为双重保障。"),
        ("Role-based access", "基于角色的访问",
         "@login_required(role='student'|'teacher') on routes.",
         "路由上使用角色装饰器限制访问。"),
        ("Teacher course ownership", "教师仅管理自己的课",
         "WHERE teacher_id = current user in SELECT/UPDATE/DELETE.",
         "SQL 条件限定 teacher_id 为当前登录教师。"),
        ("Capacity edit rule", "编辑容量规则",
         "Cannot set max_capacity below current enrolled count.",
         "不能把容量改成小于当前已选人数。"),
    ]
    for en, zh, end, zhd in rules:
        add_h3(doc, en, zh)
        add_para(doc, end, zhd)

    doc.add_page_break()

    # --- 9 Features step by step ---
    add_h1(doc, "9. Features — Step by Step", "9. 功能实现 — 逐步说明")

    features = [
        ("9.1 Student browse courses", "9.1 学生浏览课程",
         "Route: GET /student/courses\nSQL JOIN courses + teacher name; subquery COUNT for enrolled_count; EXISTS for is_enrolled.\nTemplate: student/courses.html shows cards with Enroll / Drop / Full.",
         "路由 GET /student/courses\nSQL 联结课程与教师名；子查询统计人数；EXISTS 判断是否已选。\n模板 courses.html 显示卡片与选课/退课/满员按钮。"),
        ("9.2 Student enroll", "9.2 学生选课",
         "Route: POST /student/enroll/<course_id>\nChecks: course exists, not duplicate, not full → INSERT enrollments.\nSyntax: methods=['POST'] only (not GET) for safety.",
         "路由 POST /student/enroll/<id>\n校验：课程存在、未重复、未满员 → INSERT。\n仅用 POST 方法更安全（避免误触链接选课）。"),
        ("9.3 Student drop course", "9.3 学生退课",
         "Route: POST /student/drop/<course_id>\nDELETE FROM enrollments WHERE student_id AND course_id.\nOptional ?next=browse to return to browse page.",
         "路由 POST 退课，DELETE 选课记录。\n可用 ?next=browse 返回浏览页。"),
        ("9.4 Course full test (student3)", "9.4 满员测试（student3）",
         "CS101 max_capacity=2. student1 and student2 enroll. student3 clicks Try Enroll → server flash: 'CS101 is full (2/2 seats)...'.",
         "CS101 容量为 2。student1、2 选课后，student3 尝试选课显示英文满员提示。"),
        ("9.5 Teacher add course", "9.5 教师添加课程",
         "Route: /teacher/courses/new GET shows form, POST inserts with teacher_id=current user.\ncode.strip().upper() normalizes course code.",
         "GET 显示表单，POST 插入并绑定 teacher_id。\ncode 转大写统一格式。"),
        ("9.6 Teacher edit/delete course", "9.6 教师编辑/删除课程",
         "Edit: UPDATE title, description, max_capacity.\nDelete: DELETE courses WHERE id AND teacher_id; CASCADE removes enrollments.",
         "编辑：UPDATE 课程信息。\n删除：仅删自己的课；级联删除选课记录。"),
        ("9.7 Teacher remove student", "9.7 教师移除学生",
         "POST /teacher/enrollments/<id>/remove\nVerifies course.teacher_id matches logged-in teacher.",
         "POST 移除；校验课程属于当前教师。"),
        ("9.8 Teacher assign grade", "9.8 教师录入成绩",
         "POST /teacher/grades/<enrollment_id> with form field grade.\nUPDATE enrollments SET grade = ?",
         "POST 提交成绩字段，UPDATE enrollments 表。"),
    ]
    for en, zh, body_en, body_zh in features:
        add_h2(doc, en, zh)
        add_para(doc, body_en, body_zh)

    # --- 10 Frontend ---
    add_h1(doc, "10. Frontend (HTML, CSS, JavaScript)", "10. 前端（HTML、CSS、JavaScript）")
    add_h2(doc, "10.1 Jinja2 templates", "10.1 Jinja2 模板语法")
    add_bullet(doc, "{% extends 'base.html' %} — inherit layout", "{% extends %} — 继承基础布局")
    add_bullet(doc, "{% block content %} — child fills section", "{% block content %} — 子页面填充区域")
    add_bullet(doc, "{{ variable }} — print server data", "{{ 变量 }} — 输出后端数据")
    add_bullet(doc, "{% for x in list %} — loop courses/students", "{% for %} — 循环课程/学生")
    add_bullet(doc, "{% if session.get('role') == 'student' %} — role menu", "{% if %} — 按角色显示菜单")
    add_bullet(doc, "url_for('route_name') — build correct URL", "url_for — 生成正确 URL")

    add_h2(doc, "10.2 Forms", "10.2 表单")
    add_para(
        doc,
        "<form method=\"post\" action=\"{{ url_for('enroll_course', course_id=course.id) }}\"> submits data to Flask. CSRF not implemented in this demo (noted for thesis limitation).",
        "form method=post 提交到 Flask。本演示未实现 CSRF（可在论文中作为改进点说明）。",
    )

    add_h2(doc, "10.3 CSS (style.css)", "10.3 样式表")
    add_para(
        doc,
        ":root CSS variables define colors. .card-grid uses CSS Grid. .flash-success/danger style messages. @media for mobile.",
        ":root 定义主题色。.card-grid 用网格布局。.flash-* 样式化提示。 @media 适配手机。",
    )

    add_h2(doc, "10.4 JavaScript (app.js)", "10.4 JavaScript")
    add_para(
        doc,
        "DOMContentLoaded waits for page load. data-confirm + confirm() prevents accidental enroll/delete. Flash messages fade after 5 seconds.",
        "DOMContentLoaded 等待页面加载。data-confirm 防止误操作。提示条 5 秒后淡出。",
    )

    doc.add_page_break()

    # --- 11 Problems ---
    add_h1(doc, "11. Problems Encountered & Solutions", "11. 遇到的问题与解决方案")
    problems = [
        ("GitHub repo name with spaces", "仓库名含空格",
         "gh repo create failed: use full-stack-study-system not 'Full-Stack Study System'.",
         "gh 报错：仓库名用连字符，不能有空格。"),
        ("Empty-looking GitHub page", "GitHub 页面看似为空",
         "Code was under account qwerscd; must expand templates/ folder; README updated with tree.",
         "代码在 qwerscd 账号下；需展开文件夹查看。"),
        ("Course full UI only showed badge", "满课只显示徽章无提示",
         "Added Try Enroll button + server flash message in English.",
         "增加 Try Enroll 按钮与服务端英文 flash 提示。"),
        ("Missing student3 for full test", "缺少 student3 测满员",
         "Added ensure_demo_users() to insert student3 without resetting DB.",
         "ensure_demo_users 补账号而不清空数据库。"),
        ("No drop course initially", "最初没有退课",
         "Added POST /student/drop/<course_id> and Drop buttons.",
         "新增退课路由与按钮。"),
        ("Teacher delete/remove not implemented", "教师删课/移学生未实现",
         "Added teacher_course_delete and teacher_remove_student routes.",
         "新增删除课程与移除学生路由。"),
    ]
    for en, zh, sol_en, sol_zh in problems:
        add_h3(doc, en, zh)
        add_para(doc, sol_en, sol_zh)

    # --- 12 Operations ---
    add_h1(doc, "12. How to Run & Operate", "12. 如何运行与操作")
    add_code(
        doc,
        """cd /Users/chenjunshuo/Projects/study-system
python3 -m venv venv
source venv/bin/activate
pip install -r requirements.txt
python app.py
# Open http://127.0.0.1:5000""",
    )
    add_h2(doc, "12.1 Demo accounts", "12.1 演示账号")
    add_code(
        doc,
        """Student: student1 / student123  (also student2, student3)
Teacher: teacher1 / teacher123""",
    )
    add_h2(doc, "12.2 GitHub upload", "12.2 上传 GitHub")
    add_code(
        doc,
        """gh auth login
cd study-system
gh repo create full-stack-study-system --public --source=. --remote=origin --push""",
    )

    # --- 13 Testing ---
    add_h1(doc, "13. Test Checklist for Thesis", "13. 论文测试清单")
    tests = [
        "Register new student → login",
        "Register teacher → login → create course",
        "Student enroll → appears in My Courses",
        "Student drop → removed from list",
        "Two students fill CS101 → third gets full message",
        "Teacher edit capacity, assign grade",
        "Teacher remove student, delete course",
        "Student cannot open /teacher/courses (permission)",
        "Teacher cannot open /student/courses (permission)",
    ]
    for t in tests:
        add_bullet(doc, t, t)

    # --- 14 Route table ---
    add_h1(doc, "14. Complete Route Table", "14. 完整路由表")
    routes = [
        ("GET /", "index", "Redirect login or dashboard", "跳转登录或仪表盘"),
        ("GET,POST /login", "login", "Login form", "登录"),
        ("GET,POST /register", "register", "Register", "注册"),
        ("GET /logout", "logout", "Logout", "退出"),
        ("GET /dashboard", "dashboard", "Role redirect", "按角色跳转"),
        ("GET /student/courses", "student_courses", "Browse courses", "浏览课程"),
        ("GET /student/my-courses", "student_my_courses", "My courses", "我的课程"),
        ("GET /student/grades", "student_grades", "Grades", "成绩"),
        ("POST /student/enroll/<id>", "enroll_course", "Enroll", "选课"),
        ("POST /student/drop/<id>", "drop_course", "Drop", "退课"),
        ("GET /teacher/courses", "teacher_courses", "Teacher list", "教师课程列表"),
        ("GET,POST /teacher/courses/new", "teacher_course_new", "Add course", "新建课程"),
        ("GET,POST /teacher/courses/<id>/edit", "teacher_course_edit", "Edit", "编辑"),
        ("POST /teacher/courses/<id>/delete", "teacher_course_delete", "Delete", "删除"),
        ("GET /teacher/courses/<id>/students", "teacher_course_students", "Roster", "学生名单"),
        ("POST /teacher/enrollments/<id>/remove", "teacher_remove_student", "Remove student", "移除学生"),
        ("POST /teacher/grades/<id>", "update_grade", "Set grade", "录成绩"),
    ]
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "URL"
    hdr[1].text = "Function"
    hdr[2].text = "EN"
    hdr[3].text = "ZH"
    for url, fn, en, zh in routes:
        row = table.add_row().cells
        row[0].text = url
        row[1].text = fn
        row[2].text = en
        row[3].text = zh

    # --- 15 Template files ---
    add_h1(doc, "15. Every Template File Explained", "15. 每个模板文件说明")
    templates = [
        ("base.html", "Layout shell: header nav by role, flash loop, {% block content %}, footer, loads style.css and app.js.", "布局壳：按角色导航、循环显示 flash、子页面 content 块、页脚、引入 CSS/JS。"),
        ("login.html", "POST form fields username, password. Demo accounts in <details>.", "登录表单 username/password，演示账号说明。"),
        ("register.html", "<select name=role> student/teacher. Validates on server.", "角色下拉，服务端校验。"),
        ("student/courses.html", "Loops courses; is_enrolled → Drop; full → full-notice + Try Enroll; else Enroll.", "循环课程；已选显示退课；满员显示说明与 Try Enroll。"),
        ("student/my_courses.html", "Table with Drop button per row; course_id for drop URL.", "表格每行退课按钮，传 course_id。"),
        ("student/grades.html", "Shows grade or Pending / Not yet graded badge.", "显示成绩或待评分状态。"),
        ("teacher/courses.html", "Table: Students, Edit, Delete (POST with data-confirm).", "表格：学生名单、编辑、删除（带确认）。"),
        ("teacher/course_form.html", "If course: edit mode (code readonly). Else: new course with code field.", "有 course 对象为编辑；否则新建含 code 字段。"),
        ("teacher/students.html", "Grade form + Remove form per student row.", "每行：成绩表单 + 移除按钮。"),
    ]
    for name, en, zh in templates:
        add_h3(doc, name, name)
        add_para(doc, en, zh)

    add_h1(doc, "16. Key Code Walkthrough (app.py)", "16. 核心代码走读（app.py）")
    add_h3(doc, "login_required — decorator pattern", "装饰器模式")
    add_code(
        doc,
        """def login_required(role=None):
    def decorator(view):
        @wraps(view)
        def wrapped(*args, **kwargs):
            ...
            return view(*args, **kwargs)
        return wrapped
    return decorator""",
    )
    add_para(
        doc,
        "@wraps(view) keeps the original function name for Flask debugging. Three nested functions implement a parameterized decorator.",
        "@wraps 保留原函数名。三层嵌套函数实现可传参的装饰器。",
    )

    add_h3(doc, "enroll_course — business logic order", "选课逻辑顺序")
    add_code(
        doc,
        """1. SELECT course by id
2. SELECT existing enrollment → duplicate check
3. SELECT COUNT(*) → capacity check
4. INSERT enrollment
5. flash + redirect""",
    )

    add_h3(doc, "URL dynamic parts", "URL 动态参数")
    add_para(
        doc,
        "<int:course_id> in route converts URL segment to integer and passes as course_id argument.",
        "路由中 <int:course_id> 将 URL 段转为整数并作为参数传入函数。",
    )

    add_h1(doc, "17. Security Notes for Thesis", "17. 论文中的安全说明")
    add_bullet(doc, "Passwords: Werkzeug hashing (pbkdf2/scrypt).", "密码：Werkzeug 哈希存储。")
    add_bullet(doc, "SQL: parameterized ? placeholders.", "SQL：参数化查询防注入。")
    add_bullet(doc, "Session: secret_key signs cookie.", "Session：secret_key 签名 Cookie。")
    add_bullet(doc, "Authorization: role decorator + teacher_id in SQL.", "授权：角色装饰器 + SQL 教师 ID 校验。")
    add_bullet(doc, "Limitation: no CSRF token in forms (mention in thesis).", "局限：表单未加 CSRF（论文中可写改进方向）。")

    # --- 18 Conclusion ---
    add_h1(doc, "18. Conclusion for Thesis", "18. 论文总结要点")
    add_para(
        doc,
        "This project demonstrates full-stack web development with clear separation of concerns, relational database design, password hashing, session-based authentication, role-based authorization, and user-centered interface design.",
        "本项目展示了全栈 Web 开发、关系型数据库设计、密码哈希、基于会话的认证、基于角色的授权，以及以用户为中心的界面设计。",
    )
    add_para(
        doc,
        "Future improvements: CSRF protection, password reset, admin role, REST API, deployment with Gunicorn + Nginx.",
        "可改进方向：CSRF 防护、密码找回、管理员角色、REST API、Gunicorn+Nginx 部署。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
